from pathlib import Path
from typing import Dict, Any
import docx
import chardet
from .base import FileProcessor
from core.logger import setup_logger

logger = setup_logger(__name__)


class DocumentProcessor(FileProcessor):
    """Process text documents: .docx, .txt, .md, .py, etc."""
    
    async def process(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from various document formats."""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.docx':
                text = await self._process_docx(file_path)
            elif extension in ['.txt', '.md', '.py', '.js', '.java', '.cpp', '.c', '.h']:
                text = await self._process_text_file(file_path)
            else:
                text = await self._process_text_file(file_path)
            
            metadata = {
                "file_type": extension[1:],
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size
            }
            
            chunks = self.chunk_text(text)
            
            return {
                "text": text,
                "metadata": metadata,
                "chunks": chunks,
                "success": True
            }
        
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            return {
                "text": "",
                "metadata": {"error": str(e)},
                "chunks": [],
                "success": False
            }
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        doc = docx.Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    async def _process_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file with encoding detection."""
        with open(file_path, 'rb') as file:
            raw_data = file.read()
        
        detected = chardet.detect(raw_data)
        encoding = detected['encoding'] or 'utf-8'
        
        try:
            return raw_data.decode(encoding)
        except:
            return raw_data.decode('utf-8', errors='ignore')
